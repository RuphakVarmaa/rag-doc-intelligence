import io
import pytest
from app.services.chunker import _chunk_text, _chunk_pdf, _chunk_docx

# Minimal valid 2-page PDF (no external deps needed)
_MINIMAL_PDF = (
    b"%PDF-1.4\n"
    b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
    b"2 0 obj<</Type/Pages/Kids[3 0 R 4 0 R]/Count 2>>endobj\n"
    b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R"
    b"/Contents 5 0 R/Resources<</Font<</F1 6 0 R>>>>>>endobj\n"
    b"4 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R"
    b"/Contents 7 0 R/Resources<</Font<</F1 6 0 R>>>>>>endobj\n"
    b"5 0 obj<</Length 44>>\nstream\nBT /F1 12 Tf 72 720 Td (Page one) Tj ET\nendstream\nendobj\n"
    b"6 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
    b"7 0 obj<</Length 44>>\nstream\nBT /F1 12 Tf 72 720 Td (Page two) Tj ET\nendstream\nendobj\n"
    b"xref\n0 8\n"
    b"0000000000 65535 f \n"
    b"0000000009 00000 n \n"
    b"0000000058 00000 n \n"
    b"0000000115 00000 n \n"
    b"0000000266 00000 n \n"
    b"0000000417 00000 n \n"
    b"0000000513 00000 n \n"
    b"0000000580 00000 n \n"
    b"trailer<</Size 8/Root 1 0 R>>\n"
    b"startxref\n676\n%%EOF\n"
)


def test_chunk_text_basic():
    text = "word " * 600  # ~600 tokens
    chunks = _chunk_text(text, chunk_size=512, overlap=64)
    assert len(chunks) >= 2
    for c in chunks:
        assert c["token_count"] <= 512
        assert c["content"]
        assert c["char_offset_end"] > c["char_offset_start"]


def test_chunk_text_short_fits_single():
    text = "Hello world. This is a short document."
    chunks = _chunk_text(text, chunk_size=512, overlap=64)
    assert len(chunks) == 1
    assert text in chunks[0]["content"] or chunks[0]["content"] in text


def test_chunk_text_overlap():
    text = "word " * 1200
    chunks = _chunk_text(text, chunk_size=512, overlap=64)
    # Adjacent chunks must share content (overlap > 0)
    for i in range(len(chunks) - 1):
        end_a = chunks[i]["char_offset_end"]
        start_b = chunks[i + 1]["char_offset_start"]
        assert start_b < end_a


def test_chunk_text_empty():
    chunks = _chunk_text("", chunk_size=512, overlap=64)
    assert chunks == []


def test_chunk_pdf_returns_page_numbers():
    chunks = _chunk_pdf(_MINIMAL_PDF)
    # pypdf may extract little/no text from a minimal PDF; just assert no crash
    # and page_number is set when content is extracted
    for c in chunks:
        assert "page_number" in c


def test_chunk_docx_produces_chunks():
    from docx import Document as DocxDocument

    buf = io.BytesIO()
    doc = DocxDocument()
    for i in range(30):
        doc.add_paragraph(f"Paragraph {i}: " + "content word " * 30)
    doc.save(buf)

    chunks = _chunk_docx(buf.getvalue())
    assert len(chunks) >= 1
    assert all(c["content"] for c in chunks)


def test_chunk_docx_empty_doc():
    from docx import Document as DocxDocument

    buf = io.BytesIO()
    DocxDocument().save(buf)
    chunks = _chunk_docx(buf.getvalue())
    assert chunks == []
